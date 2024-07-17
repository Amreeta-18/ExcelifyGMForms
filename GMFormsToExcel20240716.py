import docx
import pandas as pd
from docx.oxml.ns import qn
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.shared import RGBColor
import re

# Load the document
doc_path = "/Users/chattera/Downloads/RP_ Session#11_28072024-2.docx"
doc = docx.Document(doc_path)

# Data storage
lines = []
data = []

# Function to check if text is a question
def isa_question(text):
    if text.lower().startswith(('will', '[before action]', '[after action]')):
        return True
    prefixes = ['formed this sub-goal as a step to their overall goal', '[before action]', '[after action]']
    lower_text = text.lower()
    for prefix in prefixes:
        if prefix in lower_text:
            return True
    return False

def is_question(text):
    return text.lower().startswith('why?')

def starts_with_yes_no_maybe(text):
    return text.lower().startswith(('yes', 'no', 'maybe'))

def is_highlighted_or_colored(run):
    if run.font.highlight_color:
        return True
    if run.font.color and run.font.color.rgb:
        return True
    return False

def extract_highlighted_text_from_table(table):
    highlighted_texts = []

    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    if is_highlighted_or_colored(run):
                        highlighted_texts.append(run.text.strip())
    return highlighted_texts

def categorize_text(row_index, col_index):
    if row_index == 4 or row_index == 6:
        if col_index == 0:
            return "Yes"
        elif col_index == 1:
            return "Maybe"
        elif col_index == 2:
            return "No"
    return None

# current_scenario = ""
# current_subgoal = ""
# current_action = ""
# current_question = ""
# current_responses = {"Yes": "", "Maybe": "", "No": ""}

# Iterate over elements in the document
for element in doc.element.body:
    if isinstance(element, CT_P):
        para = docx.text.paragraph.Paragraph(element, doc)
        text = para.text.strip()
        # print(text)
        if text and not is_question(text):
            if not text.startswith('(e.g.,'):
                lines.append([text])
    elif isinstance(element, CT_Tbl):
        table = docx.table.Table(element, doc)

        first_two_rows = [cell.text.strip() for row in table.rows[:2] for cell in row.cells]
        skip_rows = any(starts_with_yes_no_maybe(text) for text in first_two_rows)
            
        if skip_rows:
            highlighted_texts = extract_highlighted_text_from_table(table)
            row_indices_to_process = [4, 6]
                
            for row_index in row_indices_to_process:
                if row_index < len(table.rows):
                    print("row_index is", row_index)
                    row = table.rows[row_index]
                    for col_index, cell in enumerate(row.cells):
                        print("col_index is", col_index)
                        combined_text = ""
                        for run in cell.paragraphs:
                            # text = run.text.strip()
                            text = run.text
                            print(text)
                            if text:
                                combined_text += text + " "
                        # combined_text = combined_text.strip()
                        category = categorize_text(row_index, col_index)
                        if category and combined_text:
                            lines.append([category + ": " + "Facets(" + str(highlighted_texts) + "):" + combined_text])
                        elif category:
                            lines.append([category + ": -" + combined_text])
                        elif not is_question(combined_text) and not combined_text.startswith('(e.g.,'):
                            lines.append([combined_text])

        if not skip_rows:
            for row in table.rows:
                for cell in row.cells:
                    combined_text = ""
                    for run in cell.paragraphs[0].runs:
                        text = run.text.strip()
                        if text:
                            combined_text += text + " "
                    combined_text = combined_text.strip()
                    if combined_text and not is_question(combined_text):
                        if not combined_text.startswith('(e.g.,'):
                            lines.append([combined_text])

# Convert to DataFrame
df = pd.DataFrame(lines, columns=["Line"])
# Remove consecutive duplicate lines
df = df[df["Line"].shift() != df["Line"]]

# # Function to clean text
# def clean_text(text):
#     # Using regex to remove non-alphabetic characters
#     cleaned_text = re.sub(r'^[^a-zA-Z]*', '', text)
#     return cleaned_text

# # Apply clean_text function to each element in DataFrame
# df = df.applymap(clean_text)
# print(df)

# Save to Excel
output_path = "/Users/chattera/Downloads/all_lines_output.xlsx"
df.to_excel(output_path, index=False)

print(f"Data saved to {output_path}")

# Define column headers for processed questions and responses
column_headers = [
    "Scenario", 
    "Subgoal", 
    "Action", 
    "Question", 
    "Yes", 
    "Maybe", 
    "No",
    "Motivations",
    "Information Processing Style",
    "Computer Self-Efficacy",
    "Attitude Towards Risk",
    "Learn by Process vs. Tinkering",
    "None of the above"
]

# Create an empty DataFrame with the specified columns
new_df = pd.DataFrame(columns=column_headers)

# Variables to keep track of current Scenario, Subgoal, Action, Question, and Responses
current_scenario = ""
current_subgoal = ""
current_action = ""
current_question = ""
current_responses = {"Yes": "", "Maybe": "", "No": ""}

# Define headers for matching
header_list = [
    "Motivations",
    "Information Processing Style",
    "Computer Self-Efficacy",
    "Attitude Towards Risk",
    "Learn by Process vs. Tinkering",
    "None of the above"
]

# Function to initialize a row with default values
def initialize_row():
    return {header: 0 for header in header_list}

# Function to update header presence based on text
def update_headers(text):
    header_presence = initialize_row()
    for header in header_list:
        if header.lower() in text.lower():
            header_presence[header] = 1
    return header_presence

# List to accumulate rows for new DataFrame
rows = []

# Flag to track when a row has been added
row_added = False

# Iterate over rows in all lines DataFrame
for index, row in df.iterrows():
    text = row['Line'].strip()
    
    # Debugging: Print the current line being processed
    # print(f"Processing line {index}: {text}")

    if text.startswith("Scenario"):
        # print("in Scenario", text)
        current_scenario = text.split(':', 1)[-1].strip()
        current_subgoal = ""
        current_action = ""
        current_question = ""
        current_responses = {"Yes": "", "Maybe": "", "No": ""}
        row_added = False

    elif text.startswith("Subgoal"):
        # print("in subgoal", text)
        current_subgoal = text.split(':', 1)[-1].strip()
        current_action = ""
        current_question = ""
        current_responses = {"Yes": "", "Maybe": "", "No": ""}
        # print(f"Updated Subgoal: {current_subgoal}")
        row_added = False

    elif text.startswith("Action"):
        current_action = text.split(':', 1)[-1].strip()
        current_question = ""
        current_responses = {"Yes": "", "Maybe": "", "No": ""}
        # print(f"Updated Action: {current_action}")
        row_added = False

    elif text.startswith("Yes:") or text.startswith("Maybe:") or text.startswith("No:"):
        response_type = text.split(':', 1)[0].strip()
        response_text = text.split(':', 1)[-1].strip()
        
        if response_text:
            current_responses[response_type] += response_text + " "
        
        # When "No" is encountered, save the row and reset
        if response_type == "No":
            if not row_added:
                row_data = {
                    "Scenario": current_scenario,
                    "Subgoal": current_subgoal,
                    "Action": current_action,
                    "Question": current_question,
                    "Yes": current_responses["Yes"].strip(),
                    "Maybe": current_responses["Maybe"].strip(),
                    "No": current_responses["No"].strip()
                }
                row_data.update(update_headers(" ".join(current_responses.values())))
                rows.append(row_data)
                row_added = True
                current_responses = {"Yes": "", "Maybe": "", "No": ""}

    elif isa_question(text):
            current_question = text
            row_data = {
                "Scenario": current_scenario,
                "Subgoal": current_subgoal,
                "Action": current_action,
                "Question": current_question,
                "Yes": current_responses["Yes"].strip(),
                "Maybe": current_responses["Maybe"].strip(),
                "No": current_responses["No"].strip()
            }
            
            # Reset for the new question
            current_responses = {"Yes": "", "Maybe": "", "No": ""}
            row_added = False

# Convert rows list to DataFrame
new_df = pd.DataFrame(rows, columns=column_headers)

# Function to remove text between "Facet" and "):"
def remove_between_facet_to_bracket(text):
    return re.sub(r'Facet.*\):', '', text)

# Apply the function to the 'Text' column
new_df['Yes'] = new_df['Yes'].apply(remove_between_facet_to_bracket)
new_df['Maybe'] = new_df['Maybe'].apply(remove_between_facet_to_bracket)
new_df['No'] = new_df['No'].apply(remove_between_facet_to_bracket)

print(new_df)

# Save to Excel
processed_output_path = "/Users/chattera/Downloads/processed_GMevaluation_responses.xlsx"
new_df.to_excel(processed_output_path, index=False)

print(f"Processed questions and responses saved to {processed_output_path}")